from django.shortcuts import get_object_or_404, redirect, render
import openpyxl
from django.contrib.auth.models import User
import random
import string
import PyPDF2
from transformers import pipeline
from docx import Document
import warnings
from django.contrib.auth.models import Group
from django.http import HttpResponseRedirect, JsonResponse
from .models import Teacher, Course, TeacherCourse, Section, Comment

import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from django.db import IntegrityError
# Create your views here.
global_comments_to_be_sent=[]
global_comments_that_exist=[]
def send_email_to_teachers():
    # Retrieve all teachers from the database
    teachers = Teacher.objects.all()
    subject='Account creation done'
    sender_email = 'i200908@nu.edu.pk'
    sender_pass = 'Blackshield@3128'
    # Email content
    mail_content = f"Your Feedback has been uploaded to our Website.Your Credentials have ben sent to you earlier.Please head on over to check."

    # Email subject
    subject = 'Student Feedback Analysis'

    # Sender email
    sender_email = 'i200908@nu.edu.pk'

    # Loop through teachers and send email to each
    for teacher in teachers:
        message = MIMEMultipart()
        message['From'] = sender_email
        message['To'] = teacher.email
         # message['Cc'] = cc_email
        message['Subject'] = subject

        receiver_email = teacher.email
        message.attach(MIMEText(mail_content, 'plain'))
        session = smtplib.SMTP('smtp.gmail.com', 587)  # use gmail with port
        session.starttls()  # enable security
        session.login(sender_email, sender_pass)  # login with mail_id and password
        text = message.as_string()
        session.sendmail(sender_email, [receiver_email], text)
        session.quit()
        print('Mail Sent')
    print('Mails Sent to all teachers')



def upload_pdf(request):
    current_directory = os.getcwd()
    
    def open_file_dialog():
        root = tk.Tk()
        root.withdraw()
        selected_folder = filedialog.askdirectory(title="Select a folder")
        copy_folder(selected_folder, current_directory)
        root.destroy()
    
    # Start a new thread to run the tkinter code
    t = threading.Thread(target=open_file_dialog)
    t.start()
    t.join()
    send_email_to_teachers()
    final_arr=parse_pdf(request)
    for i in global_comments_to_be_sent:
        print(i)
    perform_sentiment_analysis(global_comments_to_be_sent)

    #send_email_to_teachers()

    #return HttpResponseRedirect('parsepdf')

def read_pdf_points(pdf_file_path):
    pdf_points = []
    
    with open(pdf_file_path, 'rb') as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        for page_num in range(len(pdf_reader.pages)):
            
            page = pdf_reader.pages[page_num]
            page_text = page.extract_text()
            
            # Split the extracted text into lines and store in the array
            pdf_points.extend(page_text.splitlines())
    
            
    
    return pdf_points
import re
def find_strings_with_3_hyphens(strings):
    return [s for s in strings if (s.count("-") == 4 and "Date" in s)]
def total_students_and_feedback_recieved(line):
    # Define a regex pattern to match numbers
    pattern = r'\d+'  # This pattern matches one or more digits

    # Use re.findall to extract all numbers from the line
    numbers = re.findall(pattern, line)
    total_students = 0
    feedbacks_received = 0
    # Check if we have at least two numbers (total students and feedbacks received)
    if len(numbers) >= 2:
        total_students = int(numbers[0])
        feedbacks_received = int(numbers[1])
        print(f"Total Students: {total_students}")
        print(f"Feedbacks Received: {feedbacks_received}")
    return total_students, feedbacks_received
def instructors_name(string):
    splitted=string.split("Registered")
    return    splitted[0]

def returning_just_comments(lines_array):
    element_to_find = "Comments"
    index = lines_array.index(element_to_find)
    lines_array = lines_array[index:]
    return lines_array
def merging_multi_line_review(lines_array):
    for a_line in lines_array:
        #print(a_line)
        skip=False
        do_not_merge=["Course Content and Organization","Assignment, Quizzes  Evaluation","Instructor’s quality of delivery of lectures and classroom learning environment","Learning Material (Textbook, References Books, Videos etc)","Comments" ,"Teacher Evaluation","Learning materials (Lesson Plans, Course Notes etc.) were relevant and useful","Overall Evaluation"]
        for substr in do_not_merge:
            if substr in a_line:
                print(f"Skipping: {a_line}")
                skip=True
                break
        if skip:
            continue
        if  not a_line[-1].isdigit():
            
            element_to_find = a_line
            index = lines_array.index(element_to_find)
            i=index+1
            while not lines_array[index][-1].isdigit():
                lines_array[index]=lines_array[index]+lines_array[i]
                del lines_array[i]
                if lines_array[index][-1].isdigit():
                    break
                else:
                   for do_not in do_not_merge:
                       if do_not in lines_array[index]:
                           break
                    
            
            
    final_arr=[]
    i=0
    for a_line in lines_array:
        if a_line[-1].isdigit():
            if a_line[-2].isdigit():
                final_arr.append(a_line[:-2])
            else:

                final_arr.append(a_line[:-1])
        else:
            final_arr.append(a_line)
    return final_arr

def merging_multi_line_reviews(lines_array, start_of_doc):
    # print(lines_array[0])
    #print("we ARE HERE")
    text=lines_array[0]
    words = text.split()

# Extract the second last word and last word
    semester = words[-2]
    year = words[-1]

    print("Semester:", semester)
    print("Year:", year)
    # instructor_name=[]
    # instructor_name.append("")
    
    #    # print("we getting instructors name")
    #    # print(lines_array[2])
    # print(f"line being split:{lines_array[2]}")
    # instructor_name=lines_array[2].split("Registered")

    # registered_students=instructor_name[1]
    # print("instructor Name")
    # print(instructor_name[0])
    # print("registered sudents")
    # print(registered_students)
    
        #now we open the file and add instructor anme
    #now we find out course code and allocation number
    
    # print(lines_array)
    instructor_name=instructors_name(lines_array[2])
    #search this instructor in the teacher model and get it's foreign key
    teacher=""
    try:
    # Try to find the teacher with the given instructor_name
        teacher = Teacher.objects.get(username=instructor_name)

        # Now you can access the teacher's foreign key (teacher_id)
        teacher_foreign_key = teacher.teacher_id

        # You can also access other fields of the teacher model if needed
        teacher_roll_number = teacher.teacher_roll_number

    # Use the teacher_foreign_key or other information as needed
    # ...

    except Teacher.DoesNotExist:
    # Handle the case where the teacher with the given name does not exist
        print(f"Teacher with username '{instructor_name}' does not exist.")
    result_strings = find_strings_with_3_hyphens(lines_array)
    course,section=getting_section_and_course(result_strings[0])
    course, section = getting_section_and_course(result_strings[0])
    print("////////////////////////COURSE AND SECTION///////////////////////////////")
    print(course)
    print(section)
    print(instructor_name)

    start_index = section.find("Section:")

# Extract the section substring from start_index until the end
    section = section[start_index:].strip()
    parts = course.split('-')

    # Extract course code and course name and remove leading/trailing spaces
    course_code = parts[0].strip()
    course_name = parts[1].strip()

    print("Course Code:", course_code)
    print("Course Name:", course_name)
    parts = instructor_name.split(':')

# Extract the instructor name and remove leading/trailing spaces
    instructor_name = parts[1].strip()

    print("Instructor Name:", instructor_name)
    existing_course = Course.objects.filter(course_code=course_code).first()

    if existing_course is None:
        # If the course does not exist, create a new one
        existing_course = Course(course_code=course_code, course_name=course_name)
        existing_course.save()
        print(f"Course with code '{course_code}' added successfully.")
    else:
        print(f"Course with code '{course_code}' already exists.")
    teacher = Teacher.objects.filter(first_name=instructor_name).first()

    if teacher is not None:
        # You can now access attributes and methods of the teacher object
        print("Teacher ID:", teacher.teacher_id)
        print("Teacher Roll Number:", teacher.teacher_roll_number)
    else:
        print(f"No teacher found with the name '{instructor_name}'.")
    
    try:
        teacher_course_entry = TeacherCourse.objects.filter(teacher=teacher, course=existing_course).first()

        if teacher_course_entry is None:
            # If the entry does not exist, create it
            teacher_course_entry = TeacherCourse(teacher=teacher, course=existing_course)
            teacher_course_entry.save()
            print("TeacherCourse entry added successfully.")
        else:
            print("TeacherCourse entry already exists.")
    except IntegrityError as e:
        print(f"An error occurred: {e}")


    try:
        section_entry = Section.objects.filter(teacher_course=teacher_course_entry, section_name=section, semester=semester, year=year).first()
        if(section_entry.registered_students==0):
            section_entry.registered_students



        if section_entry is None:
            # If the entry does not exist, create it
            section_entry = Section(teacher_course=teacher_course_entry, section_name=section, semester=semester, year=year)
            section_entry.save()
            print("Section entry added successfully.")
        else:
            print("Section entry already exists.")
    except IntegrityError as e:
        print(f"An error occurred: {e}")
    #ABB WE NEED TO MAKE SECTION KEE ENTRY 
    
    print(f"Entry in Section added successfully.")
    #now we need to find the teacher 
    # try:
    #     # Try to find the course with the given course_name in the Course model
    #     course_object = Course.objects.get(course_code=course)

    #     # Now you can use the course_object as needed
    #     course_id = course_object.course_id
    #     course_code=course_object.course_code
    #     # ...

    # except Course.DoesNotExist:
    #     # Handle the case where the course with the given name does not exist
    #     print(f"Course with name '{course}' does not exist.")
    # #here we have both toh register in the teacher course table about the current teachers and course
    # try:
    #     # Get the Teacher and Course objects using their IDs
    #     teacher = Teacher.objects.get(pk=teacher_foreign_key)
    #     course = Course.objects.get(pk=course_id)

    #     # Create a new TeacherCourse instance and save it
    #     new_teacher_course = TeacherCourse.objects.create(teacher=teacher, course=course)
    #     new_teacher_course.save()

    #     # Optionally, print the newly created TeacherCourse instance
    #     print(f"Created TeacherCourse: {new_teacher_course}")

    # except Teacher.DoesNotExist:
    #     print(f"Teacher with ID '{teacher_foreign_key}' does not exist.")
    # except Course.DoesNotExist:
    #     print(f"Course with ID '{course_id}' does not exist.")

    line=lines_array[2].split("Registered")
    total_students,regiesterd_students=total_students_and_feedback_recieved(line[1])
    section_entry = Section.objects.filter(teacher_course=teacher_course_entry, section_name=section, semester=semester, year=year).first()
    if(section_entry.registered_students==0):
        section_entry.registered_students=total_students
        section_entry.feedbacks_received=regiesterd_students
        section_entry.save()
    curr_file1 = "output.docx"
    final_asnwer1 = rf'C:\Users\DELL\Desktop\our FYP stuff\{curr_file1}'
    ##############HERE WE WILL CREATE THE MODELS AND SEND THEM
    arr_to_write=[instructor_name, course, section,"total registered students: "+str(total_students),"feed backs recieved: "+str(regiesterd_students)]
    #YAHAN PAR GET KEH SECTION KEH TABLE MAIN KYAA ID HAI IT WILL BE APPENDED TO EACH COMMENT LIKE 23- SECTION ID) 

    append_array_to_word(arr_to_write, final_asnwer1, final_asnwer1,True)
    lines_array = returning_just_comments(lines_array)
    
    #abb for line in lines karke agar kisi ka end digit se nhi horaha us ko check krtay jaao keh uss keh aagay kahan jaa kar number aata hai 
    final_arr=merging_multi_line_review(lines_array)
    #YAHA PAR WE CHECK IF THE COMMENT IS IN THE DATABASE
    # JO HAIN UNN KEH SATH APPEND THE 23- SECTION ID) 
    # UNN KO NIKAAL DO  
    #baaki make a list keh jitnay comments hain naa jo jaarahay hain utnay baar section kee id aajaye

    # #BAAKEE WHAT WE CAN DO IS KEH WE CAN APPEND SECTION WAGHERA SATH AND SEND IT TO GPT ASKING IT TO IGNORE THE TEXT JO SAMNAY HAI SHURU MAIN
    #AFTER IT IS SENT TO GPT IT WILL RETURN TO US KEH KYAA POSIT
    # jabb ajayegaa gpt se toh woh bhi numbered aayegaa for each comment henaa 
    # toh abb aapko pataa hai jo first comment aayaa hai waapis woh kiss section kaa hai aur konsaa comment hai and so udhar just jaa kar store krdo database main
    append_array_to_word(final_arr, final_asnwer1, final_asnwer1,False)  
    array=final_arr
    
    category=["Course Content and Organization","Assignment, Quizzes  Evaluation","Instructor’s quality of delivery of lectures and classroom learning environment","Learning Material (Textbook, References Books, Videos etc)"]   
    do_not_merge=["Course Content and Organization","Assignment, Quizzes  Evaluation","Instructor’s quality of delivery of lectures and classroom learning environment","Learning Material (Textbook, References Books, Videos etc)","Comments" ,"Teacher Evaluation","Learning materials (Lesson Plans, Course Notes etc.) were relevant and useful","Overall Evaluation"]
    cat=""
    array = [
    comment for comment in array
    if "comments" not in comment.lower() and "page 2 of" not in comment.lower()
]
    for i in range(0,len(array)):
        array[i]= array[i].strip()
        if array[i] in category:
            #idhar hee what you do is see if that comment is in the database. if it is in the database then it exists with it's emotion
            #toh phir iss ko keep putting it another array
            #and removing it from here
            #a
            
            cat=array[i]
            continue
        existing_comment = Comment.objects.filter(comments=array[i]).first()
        if existing_comment:
            array[i]="("+cat+")"+"("+str(section_entry.pk)+")"+array[i]
        # If the comment exists, remove it from 'array' and add it to 'comments_array'
              # Set the element to None (or any marker) to be removed later
            global_comments_that_exist.append(array[i])
            array[i] = None
        else:
            array[i]="("+cat+")"+"("+str(section_entry.pk)+")"+array[i]
            global_comments_to_be_sent.append(array[i])
        
        final_arr=global_comments_to_be_sent
    print(array)
        # print("mariaa")   
        #print(a_line)
    return final_arr, section

def write_array_to_word(array, output_file):
    # Create a new Word document
    doc = Document()

    # Add the array elements to the document
    for element in array:
        doc.add_paragraph(str(element))

    # Save the document
    doc.save(output_file)
                 
def getting_section_and_course(input_string):
   # print(input_string)
    split_result = input_string.split("Section")
    print(split_result)
# The first part is everything before "Section:"
    before_section = split_result[0].strip()

# Process the second part to exclude the "Date" part
    section_and_date = split_result[1].strip().split("Date")

# The second part contains everything from "Section" to "Date" (excluding the "Date" part)
    section_without_date = "Section" + section_and_date[0].strip()

# Print the results
    print("Before Section:", before_section)
    print("Section without Date:", section_without_date)
    return before_section, section_without_date
#from docx import Document
def append_array_to_word(array, input_file, output_file, headings):
    # Create a new Word document
    new_doc = Document()

    # Open the existing Word document
    doc = Document(input_file)

    # Copy paragraphs from the existing document to the new document
    # for paragraph in doc.paragraphs:
    #     new_doc.add_paragraph(paragraph.text)

    # Add the array elements to the new document
    
    category=["Course Content and Organization","Assignment, Quizzes  Evaluation","Instructor’s quality of delivery of lectures and classroom learning environment","Learning Material (Textbook, References Books, Videos etc)"]   
    do_not_merge=["Course Content and Organization","Assignment, Quizzes  Evaluation","Instructor’s quality of delivery of lectures and classroom learning environment","Learning Material (Textbook, References Books, Videos etc)","Comments" ,"Teacher Evaluation","Learning materials (Lesson Plans, Course Notes etc.) were relevant and useful","Overall Evaluation"]
    cat=""
    # for i in range(0,len(array)):
    #     if array[i] in category:
    #         cat=array[i]
    #     array[i]=cat+array[i]
        
    
    for element in array:
        if "Page 2 of" in element:
            continue
            
        para=doc.add_paragraph()
        bold_para = para.add_run(str(element))

        if headings or element in do_not_merge:
            bold_para.bold=True
        else:
            bold_para.bold=False


    # Save the new document with the appended content
    doc.save(output_file)
all_sentiments=[]
def parse_pdf(request):
    for i in range (2,8):
        curr_file = "fpa" + str(i) + ".pdf"
        #curr_file = "fpa" + str(4) + ".pdf"
        #pdf_file_path = rf'..\feedbacks\{curr_file}'
        pdf_file_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'feedbacks', curr_file))
        # with open(pdf_file_path, 'rb') as file:
        #     pdf_reader = PyPDF2.PdfReader(file)
        #     document_info = pdf_reader.metadata()

        #     print("Title:", document_info.title)
        #     print("Author:", document_info.author)
        #     print("Subject:", document_info.subject)
        #     print("Producer:", document_info.producer)
        #     print("Created Date:", document_info.createdDate)
        #     print("Number of Pages:", pdf_reader.numPages)
        points_array = read_pdf_points(pdf_file_path)
        
        final_arr,section=merging_multi_line_reviews(points_array, i)
    return render(request, 'authentication\sentences.html')
import tkinter as tk
from tkinter import filedialog
import shutil
import os
import os
import shutil
import tkinter as tk
from tkinter import filedialog
from django.http import HttpResponse
import threading
from django.contrib import messages

def copy_folder(selected_folder, current_directory):
    if selected_folder:
        # Copy the entire selected folder and its contents to the current directory
        folder_name = os.path.basename(selected_folder)
        destination = os.path.join(current_directory, folder_name)
        shutil.copytree(selected_folder, destination)
        print(f"Folder '{folder_name}' and its contents have been copied to {current_directory}")
    else:
        print("No folder selected.")

# def upload_pdf(request):
#     current_directory = os.getcwd()
    
#     def open_file_dialog():
#         root = tk.Tk()
#         root.withdraw()
#         selected_folder = filedialog.askdirectory(title="Select a folder")
#         copy_folder(selected_folder, current_directory)
#         root.destroy()
    
#     # Start a new thread to run the tkinter codeg
#     t = threading.Thread(target=open_file_dialog)
#     t.start()
#     t.join()
#     final_arr=parse_pdf(request)
#     for i in global_comments_to_be_sent:
#         print(i)
#     perform_sentiment_analysis(global_comments_to_be_sent)
    #return HttpResponseRedirect('parsepdf')
import openai
import time
# def storing_it_in_db(request):
    #now what we need to do is jese jese we read the PDF we search the teacher name in the database aur uss kee foriegn key 
    #firstly when we read input we have teacher name and we have course name and we have semester. we have all the courses table made
    # we also have teachers table made sahee
    #abb when it reads teachers
from rest_framework.decorators import api_view
from rest_framework.response import Response
# @api_view['POST']





def perform_sentiment_analysis(unique_comments):
    gl_sent=global_comments_to_be_sent
    
    for i in range(len(global_comments_to_be_sent)):
        global_comments_to_be_sent[i] = f"{i + 1}. {global_comments_to_be_sent[i]}"
    print("/////////////////////////THE ONES BEING SENT/////////////////////")
    print(global_comments_to_be_sent)
    # Set your OpenAI API key
    # openai.api_key = "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"

    # Your result_array and formatted_comments logic here...
   # print("In function")
    # Set the maximum number of tokens per request
    max_tokens_per_request = 3000

    # Initialize an empty list to store the chunks
    chunks = []

    # Initialize an empty string to store the current chunk
    current_chunk = ""
    # Iterate through each formatted comment
    
    for comment in unique_comments:
        # Check if adding the comment to the current chunk exceeds the maximum tokens
        if len(current_chunk) + len(comment) < max_tokens_per_request:
            # Add the comment to the current chunk
           # print(comment)
            current_chunk += comment
           # current_chunk += "/"
            
        else:
            # If adding the comment exceeds the limit, start a new chunk
            chunks.append(current_chunk)
            current_chunk = comment

    # Add the last chunk to the list
    chunks.append(current_chunk)
    #print(chunks)

     # Initialize an empty list to store the responses
    responses = []

    # Iterate through each chunk and send requests
    for chunk_number, chunk in enumerate(chunks):
        print(f"The chunk is:\n {chunk}")
       # print(f"The no of comments in chunk are: {len(chunk)}")
        #prompt = f"Consider the following numbered text bullets for analysis and then list your response number-wise by writing either Yes if the input text bullet has clear criticism or negative remark in it or write Compliment if it has something nice mentioned if there is neither write not sure: you must return an array that consists of yes and no at the same index as the comment JUST THE ARRAY and no text in the answer you will give\n{chunk}"
        prompt= f"Given a set of comments,in which your task is to perform sentiment analysis and categorize each comment as either positive or constructive criticism. Please follow the guidelines below:Positive Comments:If a comment expresses positive sentiments, praise, or general satisfaction, classify it as 'Positive. 'Constructive Criticism: If a comment provides constructive feedback, suggests improvements, or points out areas of concern in a positive manner, classify it as 'Constructive Criticism.' Neutral Comments: Any comment that does not explicitly contain criticism should be considered 'Positive.' Criteria for Positive Comments: Comments expressing satisfaction, appreciation, or positive sentiments. Comments not containing any form of criticism. Criteria for Constructive Criticism: Comments that suggest improvements or provide feedback or show dislike for something. Comments that point out dissatisfaction or anger of discomfort will be marked as constructive criticism .comments like fair enough would be mentioned as positive. you must return your respose like 1.Constructive criticism 2. positive. the number should match the number of the bullet. so if the chunk sento to you is '1. good2. bad15. bad' resopnse should be '1. positive 2. Constructive criticism 15. Constructive criticism'  ignore any number that is inside () ALSO THIS IS A POSITIVE XOMMENT 'The assignments and quizzes were evaluatd fairly' SOME SPECIFIC CASES:COMMENTS THAT HAVE SOMETHING ABOUT DISLIKING SOMETHING ARE ABOUT CONSTRUCTIVE CRITICISM, 'FAIR' IS POSITIVE, 'FAIR ENOUGH' IS ALSO POSITIVE \n{chunk}"
        
      
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=400,
            temperature=0.4
        )

        # Append the response content to the list
        responses.append(response['choices'][0]['message']['content'])
        # Introduce a wait statement (e.g., wait for 10 seconds between requests)
        time.sleep(10)
        # Print the current progress
        print(f"Processed chunk {chunk_number + 1}/{len(chunks)}")
        print("\n\n")

        print(responses)
    all_comments_analysis=[]
    for i in range(0,len(chunks)):
        responses_list = responses[i].split('\n')

# Extract the sentiments from each response
        sentiments = [response.split('. ')[1] for response in responses_list if response]
        for iff in sentiments:
            all_comments_analysis.append(iff)
    for j in all_comments_analysis:
        print(j)
    print(len(all_comments_analysis))
    print(gl_sent)
    gl_sent = [re.sub(r'^\d+\.\s*', '', feedback) for feedback in gl_sent]

    for i in range(0,len(gl_sent)):
        print("sss")
        # Use regular expression to extract three parts
        match = re.match(r"\((.*?)\)\((\d+)\)(.*)", gl_sent[i])
        
        if match:
            category = match.group(1).strip()
            number = match.group(2).strip()
            comment_text = match.group(3).strip()

            # Print all three parts
            print("Category:", category)
            print("Number:", number)
            print("Comment:", comment_text)
            add_comment_to_section(int(number), comment_text, category, all_comments_analysis[i])

    # Concatenate the responses into a single string
    criticised_text = '\n'.join(responses)
   # print(f"No of criticised_text : {len(criticised_text)}")
    # print(criticised_text)
def generate_random_strings( length=10):
    characters = string.ascii_letters + string.digits + string.punctuation
    random_strings = []

    random_string = ''.join(random.choice(characters) for _ in range(length))
    #random_strings.append(random_string)

    return random_string
def register_teachers(request):
    print("here we register teacher")
    excel_file = r'C:\\Users\\DELL\Desktop\\log\\login\\somehtingg.xlsx'
    workbook = openpyxl.load_workbook(excel_file)
    sheet = workbook.active
    

    # Add the array elements to the document
    
        
    # Iterate through rows and create user accounts
    for row in sheet.iter_rows(min_row=2, values_only=True):
        teacher_id=row[0]
        email = row[1]  # Assuming email is in the first column
        name=row[2]
        if Teacher.objects.filter(email=email).exists():
            messages.error(request, "Email Already Registered!!")
        else:
            username = email # Use a portion of the email as username
            password = generate_random_strings()  # Set a default password or generate one
            s=username+password
        # Create a new user
           

    # Save the document
            doc=Document()
                 
            user = Teacher.objects.create_user(username=username, email=email, password=password, first_name=name)
            user.teacher_roll_number=teacher_id
            doc.add_paragraph(s)                        
            doc.save('usernames.docx')
            print(f'User {username} created with email {email}')
            #user = User.objects.get(username=username)
            group, created = Group.objects.get_or_create(name='Instructor')
            user.groups.add(group)
            print("////////////////NEW USER BECOMES INSTRUCTOT")
    # Close the Excel file
    workbook.close()
    messages.success(request,"you have successfully registered all teachers")
    return redirect('home')

 

def register_courses():
    #kuch
    print("register courses")

def store_teacher_course(teacher,course):
    print("register teacher and courses")
from django.db.models import F
def add_comment_to_section(section_id, comment_text, category, sentiment):
    try:
        # Get the Section object based on the provided section_id
        section = Section.objects.get(pk=section_id)
        
        if(sentiment=='Positive'):
            section.number_of_positive_comments = section.number_of_positive_comments + 1
            section.save()
            if(category=='Course Content and Organization'):
                section.number_of_positive_comments_content_org = section.number_of_positive_comments_content_org + 1
                section.save()
            elif(category=='Instructor’s quality of delivery of lectures and classroom learning environment'):
                section.number_of_positive_comments_classenv = section.number_of_positive_comments_classenv  + 1
                section.save()
            elif(category=='Learning Material (Textbook, References Books, Videos etc)'):
                section.number_of_positive_comments_learning_mat = section.number_of_positive_comments_learning_mat  + 1
                section.save()
            elif(category=='Assignment, Quizzes  Evaluation'):
                section.number_of_positive_comments_ass =section.number_of_positive_comments_ass + 1
                section.save()
            # section.number_of_positive_comments= F('number_of_positive_comments_ass')+F('number_of_positive_comments_learning_mat') + F('number_of_positive_comments_classenv')+F('number_of_positive_comments_content_org')
            # section.save()
        # Create a new Comment object
        elif (sentiment== 'Constructive criticism'):
            # section.number_of_constructive_comments=section.number_of_constructive_comments+1
            section.number_of_constructive_comments = section.number_of_constructive_comments  + 1
            section.save()
            new_comment = Comment(section=section, comments=comment_text, category=category)

            # Save the new comment to the database
            new_comment.save()

       
        print(f"Comment successfully added to Section: {section_id}")
    except Section.DoesNotExist:
        print(f"Section with ID {section_id} does not exist.")
    except Exception as e:
        print(f"An error occurred: {e}")